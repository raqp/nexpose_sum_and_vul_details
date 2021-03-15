from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt, RGBColor
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx.enum.table import WD_TABLE_DIRECTION
from docx.shared import Inches
import re


class ExecutiveSummary:

    def __init__(self, config, source_file, destination_path, input_document_type, os_platform):
        if input_document_type.lower() == 'sum':
            config = config['executive_summary']
            self.mapper = {
                'IP Address': 'Component',
                'Vulnerabilities Noted per IP address': 'Vulnerabilities Noted per Component',
                'Severity Level': 'Severity Level',
                'CVSSv2 Score': 'CVSS Score',
                'Compliance Status': 'Compliance Status',
                'Exceptions, False Positives, or Compensating Controls Noted by the ASV for this Vulnerability':
                    'Exceptions, False Positives, or Compensating Controls Noted by the ASV for this Vulnerability',
                'Scan Customer Company: ': 'Scan Customer Company: ',
                'ASV Company: ': 'ASV Company: ',
                'Remediation Step': 'Remediation Step',
                'Estimated Time': 'Estimated Time'
            }
        elif input_document_type.lower() == 'vul':
            config = config['vulnerability']
            self.mapper = {
                'Scan Customer Company:': 'Scan Customer Company: ',
                'ASV Company:': 'ASV Company: ',
                'Severity': 'Severity',
                'High': 'High',
                'IP Address': 'Component',
                'Port': 'Detected Open Ports, Services/ Protocols',
                'Evidence': 'Vulnerability',
                'Compliance Status': 'Compliance Status',
                'Exceptions, False Positives, or Compensating Controls Noted by the ASV for this Vulnerability':
                    'Details',
                'Instance': ''
            }
        self.input_document_type = input_document_type
        # self.document = Document(r'C:\Users\Acer\Desktop\work\nexpose_vulnerability\NExpose\nes.docx')
        self.document = Document(source_file)
        self.destination_path = destination_path
        self.slash = "/"
        if os_platform == 'windows':
            self.slash = "\\"
        self.destination = destination_path if destination_path.endswith(self.slash) else destination_path + self.slash
        self.file_name = f'report_{datetime.now().strftime("%d-%b-%Y-%H-%M")}.docx'
        self.paragraphs = self.document.paragraphs
        self.tables = self.document.tables
        self.font_name = config['font_name']
        self.table_header_font_size = config['table_header_font_size']
        self.page_height = config['page_height']
        self.page_width = config['page_width']
        self.default_color = config['default_color']
        self.p_18_font_size = config['p_18_font_size']
        self.p_12_font_size = config['p_12_font_size']
        self.p_10_font_size = config['p_10_font_size']
        self.p_10_font_color = config['p_10_font_color']
        self.table_header_color = config['table_header_color']
        self.table_content_font_size = config['table_content_font_size']
        self.severity_level = ''
        self.cvss_v2_score = ''
        self.cve_numbers = 'N/A'

    def start(self):
        self.set_font_name()
        self.iterate()
        if self.input_document_type == 'vul':
            self.add_cells()
        self.change_tables()
        self.delete_unuseful_tables()
        if self.input_document_type == 'vul':
            self.set_table_cell_text()
        self.set_all_tables_backgrounds()
        self.save_document()

    def set_font_name(self):
        self.document.styles['Normal'].font.name = self.font_name

    @staticmethod
    def iter_block_items(parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def delete_unuseful_tables(self):
        for active_table in self.tables:
            if active_table.cell(0, 0).paragraphs[0].text == '' or \
                    active_table.cell(0, 0).paragraphs[0].text.startswith('1 Scan'):
                active_table._element.getparent().remove(active_table._element)

    @staticmethod
    def set_object_color(obj, rgb_color):
        if rgb_color:
            obj.font.color.rgb = RGBColor(*rgb_color)

    @staticmethod
    def set_paragraph_font_size(run, font_size, bold=None):
        run.font.size = Pt(font_size)
        if bold is not None:
            run.bold = bold

    def iterate(self):
        for block in self.iter_block_items(self.document):
            if isinstance(block, Paragraph):
                font_size = 9
                color = None

                if block.text.startswith('Part 2a') or block.text.startswith(
                        '2') or 'table of contents' in block.text.lower():
                    block.text = ''
                elif block.text.startswith('Part 2b'):
                    index = block.text.index('.')
                    text = f'Part 2{block.text[index:]}'
                    block.text = text
                if block.text:
                    if block.style.font.size == Pt(18):
                        if self.input_document_type == 'vul':
                            block.text = self.format_paragraph_text(block.text)
                        if 'IP Address' in block.text:
                            block.text = block.text[:block.text.index('IP')] + 'Component'
                        color = self.default_color
                        font_size = self.p_18_font_size
                    elif block.style.font.size == Pt(12):
                        if 'IP Address' in block.text:
                            block.text = block.text[:block.text.index('IP')] + 'Component'
                        color = self.default_color
                        font_size = self.p_12_font_size
                    elif block.style.font.size == Pt(10):
                        color = self.p_10_font_color
                        font_size = self.p_10_font_size
                    self.set_object_color(block.runs[0], color)
                    self.set_paragraph_font_size(block.runs[0], font_size)

    def set_table_cell_text(self):
        for table in self.tables:
            if len(table.rows[0].cells) > 2:
                if table.rows[0].cells[-1].text == 'Component':
                    severity_level = table.rows[1].cells[3].text
                    cve_number = table.rows[1].cells[4].text
                    cvss_score = table.rows[1].cells[2].text
                    if len(table.rows) > 2:
                        for row in table.rows[2:]:
                            row.cells[3].paragraphs[0].add_run().text = severity_level
                            row.cells[4].paragraphs[0].add_run().text = cve_number
                            row.cells[2].paragraphs[0].add_run().text = cvss_score

    @staticmethod
    def format_paragraph_text(item):
        text = item
        if item[0].isdigit():
            index = item.find(' ')
            text = f'{item[index + 1:]}'
        return text

    def change_tables(self):
        style = None
        for index, table in enumerate(self.tables):
            borders = ['left', 'right', 'top']
            special = False
            for cell in table.rows[0].cells:
                self.set_table_header_bg_color(cell)
                for ind, paragraph in enumerate(cell.paragraphs):
                    if ind == 0:
                        style = paragraph.style
                    paragraph.style = style
                    for run in paragraph.runs:
                        if run.text in self.mapper:
                            run.text = self.mapper[run.text]
                        self.set_paragraph_font_size(run, 9, bold=False)
                        if index != 2:
                            self.set_object_color(run, self.default_color)
            if 'Scan Customer Company' in table.rows[0].cells[0].text:
                special = True
            if index == 2 and self.input_document_type != 'vul':
                borders.append('bottom')
            self.set_table_styling(table, *borders, special=special)

    def add_cells(self):
        previous_text = ""
        current_text = ""
        previous_paragraph = ""
        current_paragraph = ""
        cve = None
        instance = ""

        for index, table in enumerate(self.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        previous_paragraph = current_paragraph
                        current_paragraph = paragraph
                        if not isinstance(previous_paragraph, str) and previous_paragraph.text == 'References':
                            cve = self.parse_hyperlinks(paragraph)
                        for run in paragraph.runs:
                            previous_text = current_text
                            current_text = run.text
                            if previous_text == 'Severity':
                                self.severity_level = current_text
                            elif previous_text == 'CVSSv2 Score':
                                self.cvss_v2_score = current_text[:current_text.index(' ')]
                            elif previous_text == 'References':
                                self.cve_numbers = cve if cve else 'N/A'
                            elif run.text.startswith('IP Address'):
                                instance = self.tables[index].rows[1].cells[2].paragraphs[0].text
                                if instance:
                                    self.tables[index].rows[1].cells[1].paragraphs[0].text += f'/{instance}'
                                self.create_new_columns(table)
                                self.swap_columns_info(table)

        for table in self.tables:
            if table.rows[0].cells[0].paragraphs[0].runs[0].text == 'IP Address':
                self.delete_columns(table, [0, 1, 4, 5])
                table.columns[0].table.table_direction = WD_TABLE_DIRECTION.RTL

    @staticmethod
    def delete_columns(table, columns):
        columns.sort(reverse=True)
        grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
        for ci in columns:
            for cell in table.column_cells(ci):
                cell._tc.getparent().remove(cell._tc)

            # Delete column reference.
            col_elem = grid[ci]
            grid.remove(col_elem)

    def swap_columns_info(self, table):
        y = {-1: 0, -2: 1, -3: 4, 2: 5}
        for index, row in enumerate(table.rows):
            for i in y:
                run = row.cells[i].paragraphs[0].add_run()
                run.text = row.cells[y[i]].paragraphs[0].text
                self.set_paragraph_font_size(run, self.table_content_font_size)

    def create_new_columns(self, table):
        headers = ['CVSS Score', 'CVE Number', 'Severity Level', 'Vulnerability', 'Evidence', 'Port',
                   'IP Address']
        headers_content = {'CVE Number': self.cve_numbers, 'CVSS Score': self.cvss_v2_score,
                           'Severity Level': self.severity_level}
        if table.rows[0].cells[0].paragraphs[0].runs[0].text.startswith('IP Address'):
            for i in range(6):
                table.add_column(Inches(1.0))
                if i < 3:
                    self.add_info_into_table(table, headers[i], headers_content[headers[i]])
        self.cve_numbers = 'N/A'
        self.cvss_v2_score = ''
        self.severity_level = ''

    def add_info_into_table(self, table, header, info=None):
        table.rows[0].cells[-1].paragraphs[0].add_run().text = header
        if info:
            run = table.rows[1].cells[-1].paragraphs[0].add_run()
            run.text = info
            self.set_paragraph_font_size(run, self.table_content_font_size)

    @staticmethod
    def parse_hyperlinks(paragraph):
        result = ''
        xml = paragraph.paragraph_format.element.xml
        xml_str = str(xml)
        wt_list = re.findall('<w:t>C[\S\s]*?</w:t>', xml_str)
        if wt_list:
            wt_list = [item[item.find('>') + 1: item.rfind('<')] for item in wt_list]
            result = ', '.join([item for item in wt_list if item.startswith('CVE')])
        return result

    @staticmethod
    def set_table_header_bg_color(cell):
        tc = cell._tc
        tbl_cell_properties = tc.get_or_add_tcPr()
        cl_shading = OxmlElement('w:shd')
        cl_shading.set(qn('w:fill'), "ffffff")
        tbl_cell_properties.append(cl_shading)
        return cell

    def set_all_tables_backgrounds(self):
        for table in self.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.set_table_header_bg_color(cell)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self.set_paragraph_font_size(run, self.table_content_font_size)

    @staticmethod
    def set_table_styling(table, *args, special=False):
        tbl = table._tbl
        cell_number = 0
        coll_count = len(table.columns)
        borders = ['left', 'right', 'top', 'bottom']
        for cell in tbl.iter_tcs():
            tc_pr = cell.tcPr
            tc_borders = OxmlElement("w:tcBorders")
            for border in args:
                side = OxmlElement(f'w:{border}')
                side.set(qn("w:val"), "nil")
                tc_borders.append(side)
            for i in set(borders).difference(args):
                side = OxmlElement(f'w:{i}')
                side.set(qn("w:val"), "single")
                if cell_number < coll_count:
                    side.set(qn("w:sz"), "12")
                    cell_number += 1
                    side.set(qn("w:color"), "4f2d7f")
                else:
                    if special:
                        side.set(qn("w:color"), "ffffff")
                    else:
                        side.set(qn("w:sz"), "5")
                        side.set(qn("w:color"), "b5b5b5")
                tc_borders.append(side)
            tc_pr.append(tc_borders)

    def save_document(self):
        self.document.save(f'{self.destination_path}{self.file_name}')
